import streamlit as st
from io import BytesIO
from docx import Document
import re

st.set_page_config(page_title="Portal Match CV", layout="wide")

st.title("Portal de Match e Adaptação de CV")

# -------------------------
# ENTRADAS
# -------------------------

empresa = st.text_input("Nome da Empresa")
descricao_vaga = st.text_area("Cole aqui a descrição completa da vaga", height=300)
upload = st.file_uploader("Upload do seu CV atual (PDF ou DOCX)", type=["pdf", "docx"])

# -------------------------
# FUNÇÕES
# -------------------------

def extrair_texto(file):
    try:
        if file.type == "application/pdf":
            from pypdf import PdfReader
            reader = PdfReader(file)
            return "".join(page.extract_text() or "" for page in reader.pages)
        else:
            doc = Document(file)
            return "\n".join(p.text for p in doc.paragraphs)
    except:
        return ""

def extrair_skills(texto):
    tech_stack = [
        "kotlin", "java", "compose", "firebase",
        "mvvm", "hilt", "retrofit", "android",
        "git", "junit", "coroutines"
    ]
    texto = texto.lower()
    return [skill for skill in tech_stack if skill in texto]

def calcular_score(skills_vaga, skills_cv):
    if not skills_vaga:
        return 0, [], []
    match = set(skills_vaga).intersection(set(skills_cv))
    faltantes = set(skills_vaga) - set(skills_cv)
    score = (len(match) / len(skills_vaga)) * 100
    return round(score, 2), list(match), list(faltantes)

def gerar_docx(texto):
    doc = Document()
    for linha in texto.split("\n"):
        doc.add_paragraph(linha)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def limpar_nome(texto):
    return re.sub(r'[^a-zA-Z0-9]', '', texto)

# -------------------------
# PROCESSAMENTO
# -------------------------

if st.button("Analisar e Adaptar"):

    if not descricao_vaga or not upload or not empresa:
        st.warning("Preencha todos os campos.")
    else:
        texto_cv = extrair_texto(upload)

        skills_vaga = extrair_skills(descricao_vaga)
        skills_cv = extrair_skills(texto_cv)

        score, match, faltantes = calcular_score(skills_vaga, skills_cv)

        st.subheader("Resultado da Análise")
        st.metric("Compatibilidade", f"{score}%")

        st.write("### Skills encontradas")
        st.write(match if match else "Nenhuma identificada")

        st.write("### Skills faltantes")
        st.write(faltantes if faltantes else "Nenhuma")

        # -------- ADAPTAÇÃO --------

        texto_adaptado = f"""Carlos Ribeiro

Resumo Profissional:
Desenvolvedor Android com experiência em {", ".join(match) if match else "tecnologias relacionadas ao ecossistema Android"}.
Experiência em arquitetura MVVM, consumo de APIs REST e boas práticas de desenvolvimento.
Perfil alinhado com os requisitos da vaga na {empresa}.

Principais Competências:
{", ".join(match) if match else "A revisar conforme requisitos da vaga."}

"""

        arquivo_docx = gerar_docx(texto_adaptado)

        nome_arquivo = f"CarlosRibeiro_AndroidDeveloper_{limpar_nome(empresa)}.docx"

        st.download_button(
            label="📥 Baixar CV Adaptado",
            data=arquivo_docx,
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
